"""从评分标准附件（PDF / DOC / DOCX）中提取纯文字。

提取出的文字会拼进 AI 批改的提示词，所以统一截断到 12000 字符，
避免撑爆模型上下文。
"""
import shutil
import subprocess
from pathlib import Path

from pypdf import PdfReader
from docx import Document

from backend.exceptions.business import CriteriaExtractionError


def _extract_pdf(filepath: str) -> str:
    """逐页提取 PDF 文字（仅支持可复制文字的 PDF，扫描件提不出内容）。"""
    try:
        reader = PdfReader(filepath)
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception as exc:
        raise CriteriaExtractionError("PDF 评分标准读取失败，请检查文件是否损坏。") from exc


def _extract_docx(filepath: str) -> str:
    """提取 DOCX 的段落和表格文字（表格每行用「 | 」连接成一行）。"""
    try:
        document = Document(filepath)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise CriteriaExtractionError("DOCX 评分标准读取失败，请检查文件是否损坏。") from exc


def _extract_doc(filepath: str) -> str:
    """提取旧版 DOC 文字：依赖 macOS 自带的 textutil，其他系统会提示转格式。"""
    converter = shutil.which("textutil")
    if not converter:
        raise CriteriaExtractionError("当前环境暂不支持读取 DOC 文件，请转换为 PDF 或 DOCX 后再上传。")
    try:
        # textutil 是外部命令，必须设置 timeout，避免异常 DOC 文件让请求长期挂起。
        result = subprocess.run(
            [converter, "-convert", "txt", "-stdout", filepath],
            check=True,
            capture_output=True,
            text=True,
            timeout=30,
        )
    except (OSError, subprocess.SubprocessError) as exc:
        raise CriteriaExtractionError("DOC 评分标准读取失败，请转换为 PDF 或 DOCX 后再上传。") from exc
    return result.stdout.strip()


def extract_criteria_text(filepath: str) -> str:
    """按扩展名分发提取，返回截断后的文字（本模块唯一对外入口）。"""
    suffix = Path(filepath).suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(filepath)
    elif suffix == ".docx":
        text = _extract_docx(filepath)
    elif suffix == ".doc":
        text = _extract_doc(filepath)
    else:
        raise CriteriaExtractionError("评分标准附件格式不受支持。")

    if not text:
        raise CriteriaExtractionError("没有从评分标准附件中提取到文字，请上传可复制文字的 PDF 或 Word 文件。")
    # 模型提示词还会拼接 OCR 文本，评分标准先截断，给作业内容留下上下文空间。
    return text[:12000]
